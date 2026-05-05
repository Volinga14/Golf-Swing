from __future__ import annotations

from pathlib import Path
import sys

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Golf Swing Ai App Documentacion V0 3 MVP.docx"
TABLE_HELPERS = Path(
    r"C:\Users\v.borges\.codex\plugins\cache\openai-primary-runtime\documents\26.430.10722\skills\documents\scripts"
)
sys.path.append(str(TABLE_HELPERS))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


ACCENT = "1F5D49"
INK = "14201E"
MUTED = "66736F"
GOLD = "C99432"
PAPER = "F8F4E8"
LINE = "D8D5CA"
BLUE = "355C7D"
CORAL = "B45F4D"


def main() -> None:
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_document(doc)
    build_cover(doc)
    build_body(doc)
    doc.save(OUT)
    print(OUT)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.88)
    section.right_margin = Inches(0.88)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color in [
        ("Title", 24, ACCENT),
        ("Subtitle", 12, MUTED),
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 13, INK),
        ("Heading 3", 11, BLUE),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = style_name in {"Title", "Heading 1", "Heading 2", "Heading 3"}

    styles["Heading 1"].paragraph_format.space_before = Pt(13)
    styles["Heading 1"].paragraph_format.space_after = Pt(7)
    styles["Heading 2"].paragraph_format.space_before = Pt(9)
    styles["Heading 2"].paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.text = "Swing Lab AI MVP v0.3"
    header.style = doc.styles["Normal"]
    header.runs[0].font.size = Pt(8.5)
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    add_bottom_border(header, LINE)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Página ")
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)
    add_field(footer, "PAGE")
    footer.add_run(" de ")
    add_field(footer, "NUMPAGES")


def build_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(14)
    run = p.add_run("Golf Swing AI App")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor.from_string(ACCENT)

    subtitle = doc.add_paragraph("Documentación técnica y MVP implementado v0.3", style="Subtitle")
    subtitle.paragraph_format.space_after = Pt(18)

    meta = [
        ("Fecha", "2026-05-04"),
        ("Entregable", "PWA local-first funcional + documentación actualizada"),
        ("Objetivo", "Convertir un vídeo de swing en revisión visual, métricas simples y coaching accionable."),
        ("Estado", "MVP Sprint 1 ampliado; preparado para integrar MediaPipe en Sprint 2."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    widths = exact_widths([2100, 7260])
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=0)
    mark_header_row(table.rows[0])
    shade_cell(table.rows[0].cells[0], PAPER)
    shade_cell(table.rows[0].cells[1], PAPER)
    set_cell(table.rows[0].cells[0], "Campo", bold=True)
    set_cell(table.rows[0].cells[1], "Valor", bold=True)
    for label, value in meta:
        cells = table.add_row().cells
        set_cell(cells[0], label, bold=True)
        set_cell(cells[1], value)
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=0)

    add_callout(
        doc,
        "Principio del producto",
        "La app debe ser entretenida y fácil de usar sin perder rigor: mide solo lo que puede explicar, separa observaciones de estimaciones y convierte cada análisis en una prioridad práctica.",
        color=ACCENT,
    )

    doc.add_page_break()


def build_body(doc: Document) -> None:
    add_heading(doc, "1. Resumen Ejecutivo")
    for text in [
        "Se ha construido un MVP de interfaz completa para analizar swings desde el navegador, sin backend y sin subir vídeos a servidores. El foco es que el usuario pueda cargar un vídeo, revisarlo frame a frame, marcar eventos clave, dibujar referencias, recibir un reporte simple y guardar la sesión localmente.",
        "El MVP no pretende simular un launch monitor ni declarar métricas biomecánicas que aún no están detectadas por visión real. La versión entregada es la base Sprint 1: visor, datos, reporte, exportaciones y estructura preparada para MediaPipe Pose Landmarker Web.",
    ]:
        doc.add_paragraph(text)

    add_heading(doc, "2. Qué Se Ha Implementado")
    rows = [
        ("Subida de vídeo", "Carga local de MP4/MOV/WebM y reproducción en navegador.", "Completo"),
        ("Vista FO / DTL", "Selección de vista para cambiar guías y reglas de recomendación.", "Completo"),
        ("Orientación", "Reconoce vídeo vertical u horizontal y ajusta el visor.", "Completo"),
        ("Frame-by-frame", "Slider, botones, atajos, cámara lenta y pantalla completa.", "Completo"),
        ("Fases del swing", "Detección automática por movimiento, salto al frame y corrección manual.", "Completo"),
        ("Overlays", "Guía ajustable, fases ocultables, grid, líneas y ángulos manuales.", "Completo"),
        ("Métricas MVP", "Capture score automático, tempo ratio y métricas revisables con evidencia.", "Completo"),
        ("Recomendaciones", "Motor de reglas con prioridad, recomendaciones extra y explicación de resultados.", "Completo"),
        ("Bola y golpe", "Vista separada para disfrutar el golpe y dibujar/sugerir trayectoria de bola.", "Completo"),
        ("Historial local", "Guardado de sesiones en IndexedDB.", "Completo"),
        ("Exportaciones", "JSON, CSV de métricas y PNG del frame actual.", "Completo"),
        ("Guardrails de confianza", "Sin vídeo no hay score; las fases deben estar en orden cronológico.", "Completo"),
        ("Pruebas", "Smoke tests, validación de rutas, lógica de recomendaciones y captura headless.", "Completo"),
        ("PWA", "Manifest y service worker para ejecución local por HTTP.", "Base lista"),
    ]
    add_table(doc, ["Área", "Descripción", "Estado"], rows, [1.2, 3.4, 1.0])

    add_heading(doc, "3. Flujo De Uso")
    steps = [
        "Subir un vídeo de swing grabado con móvil.",
        "Elegir vista: down-the-line, face-on o no segura.",
        "Completar el capture score para indicar calidad de grabación.",
        "Revisar el swing con el slider o los controles frame-by-frame.",
        "Marcar address, top, impact y finish. El MVP propone una primera posición por porcentaje del vídeo.",
        "Dibujar líneas o ángulos si se quiere comparar plano, postura o posición de manos.",
        "Leer el reporte: score general, confianza, prioridad, evidencia y drill.",
        "Guardar la sesión o exportar JSON/CSV/PNG.",
    ]
    for step in steps:
        doc.add_paragraph(step, style="List Number")

    add_heading(doc, "4. Arquitectura Entregada")
    rows = [
        ("Frontend", "HTML, CSS y JavaScript modular sin dependencias externas."),
        ("Vídeo", "HTML5 video + canvas overlay superpuesto."),
        ("Estado", "Objeto de sesión en memoria con eventos, métricas, capture score y dibujos."),
        ("Storage", "IndexedDB para historial local de sesiones."),
        ("Export", "Descarga de JSON, CSV y PNG desde el propio navegador."),
        ("PWA", "Manifest + service worker. Funciona mejor servido por http://localhost."),
    ]
    add_table(doc, ["Componente", "Implementación MVP"], rows, [1.35, 4.25])

    add_heading(doc, "5. Mapa De Archivos")
    rows = [
        ("app/index.html", "Estructura principal de la experiencia."),
        ("app/styles/main.css", "Sistema visual responsive y layout de la app."),
        ("app/src/main.js", "Estado principal, eventos de UI, render de reporte e historial."),
        ("app/src/video-player.js", "Carga de vídeo, seek por frame y controles de transporte."),
        ("app/src/overlays.js", "Canvas: guías, esqueleto visual, fases, líneas y ángulos."),
        ("app/src/metrics.js", "Cálculo de métricas del MVP."),
        ("app/src/recommendations.js", "Reglas de coaching interpretables."),
        ("app/src/storage.js", "Persistencia local con IndexedDB."),
        ("app/src/export.js", "Exportación JSON, CSV y PNG."),
        ("app/assets/swing-guide.svg", "Visual inicial para la zona de captura."),
    ]
    add_table(doc, ["Archivo", "Rol"], rows, [1.8, 3.8])

    add_heading(doc, "6. Métricas MVP v0.1")
    rows = [
        ("Capture score", "Observada por checklist", "Encuadre, luz, estabilidad, visibilidad de cuerpo, bola, palo y FPS.", "Alta si el usuario responde bien."),
        ("Tempo ratio", "Calculada", "Tiempo address-top dividido entre top-impact.", "Media; depende del marcado correcto."),
        ("Head stability", "Manual asistida", "Proxy de movimiento de cabeza, especialmente útil en FO.", "Media-baja hasta MediaPipe."),
        ("Posture retention", "Manual asistida", "Proxy de conservación de postura y espacio antes de impacto.", "Media-baja hasta MediaPipe."),
        ("Hand path", "Manual asistida", "Proxy de ruta de manos y transición en DTL.", "Media-baja hasta tracking real."),
        ("Finish balance", "Manual asistida", "Estabilidad y final completo.", "Media."),
    ]
    add_table(doc, ["Métrica", "Tipo", "Qué mide", "Confianza"], rows, [1.25, 1.05, 2.3, 1.0])

    add_heading(doc, "7. Motor De Recomendaciones")
    doc.add_paragraph(
        "El MVP usa reglas interpretables. La app evita dar veinte correcciones: el reporte elige una prioridad principal, muestra evidencia, propone un drill y sugiere qué métrica mirar en la siguiente sesión."
    )
    rows = [
        ("Capture score bajo", "Repetir captura", "La grabación no permite un análisis fiable."),
        ("Tempo fuera de rango", "3:1 tempo drill", "Backswing y downswing quedan lejos de una referencia inicial cercana a 3:1."),
        ("FO + cabeza baja", "Pivot sobre eje central", "Posible sway lateral."),
        ("DTL + postura baja", "Chair drill", "Posible early extension."),
        ("DTL + ruta manos baja", "Pump drill bajo plano", "Posible transición por encima del plano."),
        ("Finish bajo", "Hold finish", "Equilibrio final poco estable."),
    ]
    add_table(doc, ["Regla", "Drill", "Motivo"], rows, [1.5, 1.5, 2.6])

    add_heading(doc, "8. Modelo De Datos Exportado")
    doc.add_paragraph(
        "El JSON exportado deja preparada la estructura para crecer hacia landmarks, métricas por frame y modelos automáticos."
    )
    rows = [
        ("session id / createdAt", "Identificación local de sesión."),
        ("video", "Nombre, duración, FPS y frames totales."),
        ("club / viewType / ballResult", "Contexto de captura y golpe."),
        ("events", "Address, top, impact y finish con frame, timestamp y confianza."),
        ("manualMetrics", "Controles rápidos usados para alimentar el reporte."),
        ("metrics", "Resultados calculados por el MVP."),
        ("recommendations", "Reglas activadas, prioridad, drill y evidencia."),
        ("overlayDrawings", "Líneas y ángulos dibujados por el usuario."),
    ]
    add_table(doc, ["Campo", "Contenido"], rows, [1.8, 3.8])

    add_heading(doc, "9. Límites Declarados Del MVP")
    limits = [
        "No hay aún MediaPipe ni landmarks reales. Se ha retirado el esqueleto falso; la guía solo sirve para encuadre y comparación.",
        "No se calcula spin rate, smash factor, carry exacto, attack angle real ni face angle.",
        "El vídeo queda local; el historial guarda metadatos y análisis, no el archivo de vídeo completo.",
        "Las recomendaciones son reglas de producto iniciales basadas en timing, movimiento y métricas revisables, no diagnóstico profesional cerrado.",
        "La trayectoria de bola es una vista visual separada: ayuda a disfrutar y registrar resultado, pero no entra en el análisis técnico del swing.",
        "El objetivo actual es validar flujo, comprensión del reporte y estructura de datos.",
    ]
    for limit in limits:
        doc.add_paragraph(limit, style="List Bullet")

    add_callout(
        doc,
        "Regla de confianza",
        "La app debe mostrar cuándo una conclusión es observada, estimada o inferida. Esto protege la confianza del usuario y evita prometer precisión que el vídeo móvil no puede dar todavía.",
        color=CORAL,
    )

    add_heading(doc, "10. Pruebas Y Criterios De Calidad")
    doc.add_paragraph(
        "El MVP incluye pruebas automatizadas para proteger la base del producto antes de añadir MediaPipe o lógica más pesada."
    )
    rows = [
        ("Smoke test", "Valida assets, ids críticos, rutas referenciadas y cache del service worker."),
        ("Lógica de métricas", "Comprueba tempo, capture score, no-vídeo, fases incompletas y fases fuera de orden."),
        ("Recomendaciones", "Verifica que la prioridad y el drill salen de reglas interpretables."),
        ("Servidor temporal", "Sirve la PWA por HTTP y confirma que los archivos principales responden."),
        ("Headless browser", "Abre la app con Microsoft Edge headless, confirma DOM post-JS y genera captura PNG."),
        ("Documento", "Auditoría estructural: tablas con ancho fijo, encabezados de tabla y cero findings de accesibilidad estructural."),
    ]
    add_table(doc, ["Prueba", "Qué protege"], rows, [1.7, 3.9])

    add_heading(doc, "11. Siguiente Sprint: Pose Estimation")
    rows = [
        ("Integrar MediaPipe Pose Landmarker Web", "Extraer 33 landmarks por frame en navegador."),
        ("Dibujar skeleton real", "Sustituir guía visual por overlay basado en detección."),
        ("Suavizado temporal", "Reducir jitter de landmarks en swings rápidos."),
        ("CSV por frame", "Exportar landmarks y métricas base."),
        ("Confianza por métrica", "Desactivar recomendaciones cuando la pose sea mala."),
        ("Corrección manual", "Permitir que el usuario ajuste eventos cuando el modelo falle."),
    ]
    add_table(doc, ["Tarea", "Resultado esperado"], rows, [2.05, 3.55])

    add_heading(doc, "12. Roadmap Actualizado")
    rows = [
        ("Sprint 1", "PWA usable: vídeo, fases, overlays, métricas rápidas, recomendaciones, historial y export.", "Entregado"),
        ("Sprint 2", "MediaPipe, skeleton real, landmarks, CSV por frame.", "Siguiente"),
        ("Sprint 3", "Métricas biomecánicas FO/DTL más robustas y score de confianza.", "Planificado"),
        ("Sprint 4", "Segmentación automática inicial con heurísticas de manos/landmarks.", "Planificado"),
        ("Sprint 5", "Ball tracker nivel 1 y selección manual asistida.", "Futuro"),
        ("Sprint 6", "Club tracker nivel 1 y plano del palo.", "Futuro"),
        ("Sprint 7", "Comparativa antes/después e historial avanzado.", "Futuro"),
    ]
    add_table(doc, ["Fase", "Objetivo", "Estado"], rows, [1.0, 3.7, 0.9])

    add_heading(doc, "13. Definición De Éxito Del MVP")
    success = [
        "Un usuario puede analizar su primer vídeo sin crear cuenta ni subirlo a un servidor.",
        "El reporte se entiende en menos de 30 segundos.",
        "La app produce una recomendación práctica y una métrica a vigilar.",
        "El usuario puede guardar una sesión y exportar datos.",
        "El alcance técnico queda claro: herramienta de revisión y preparación para IA, no launch monitor.",
    ]
    for item in success:
        doc.add_paragraph(item, style="List Bullet")

    add_heading(doc, "14. Decisión Técnica Reafirmada")
    doc.add_paragraph(
        "Se confirma la opción A del documento v0.1: PWA pura primero. La app ya funciona como interfaz de usuario y laboratorio visual. Python/OpenCV queda como vía paralela para experimentar métricas o modelos antes de migrarlos al navegador o a un backend."
    )


def add_heading(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="Heading 1")


def add_table(doc: Document, headers: list[str], rows: list[tuple[str, ...]], weights: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = 0
    widths = column_widths_from_weights(weights, total_width_dxa=9360)
    mark_header_row(table.rows[0])
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, PAPER)
        set_cell(cell, header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell(cells[idx], value)
    apply_table_geometry(table, widths, table_width_dxa=sum(widths), indent_dxa=0)
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
    doc.add_paragraph()


def set_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(9.2)
    run.font.color.rgb = RGBColor.from_string(INK if bold else "25302E")


def add_callout(doc: Document, label: str, body: str, color: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(8)
    p.paragraph_format.right_indent = Pt(8)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(10)
    shade_paragraph(p, "F8F4E8")
    add_left_border(p, color)
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string(color)
    run = p.add_run("\n" + body)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(10.2)
    run.font.color.rgb = RGBColor.from_string(INK)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def shade_paragraph(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def add_left_border(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "12")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), color)
    p_bdr.append(left)


def mark_header_row(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = tr_pr.find(qn("w:tblHeader"))
    if tbl_header is None:
        tbl_header = OxmlElement("w:tblHeader")
        tr_pr.append(tbl_header)
    tbl_header.set(qn("w:val"), "true")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_bottom_border(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_field(paragraph, field: str) -> None:
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor.from_string(MUTED)


def exact_widths(widths: list[int]) -> list[int]:
    return widths


if __name__ == "__main__":
    main()
